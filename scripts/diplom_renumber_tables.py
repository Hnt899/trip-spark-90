# -*- coding: utf-8 -*-
"""
Сквозная нумерация таблиц в ПЗ (Word): 1, 2, 3 … вместо 1.2, 2.1, 3.1 …;
подписи «Таблица N – …»; вставка подписей для таблиц 2 и 3 перед таблицами.
Требует: pip install python-docx

Запуск: python scripts/diplom_renumber_tables.py [путь_к.docx]
По умолчанию: _diplom_6pz_vps.docx → перезапись и копия на рабочий стол.
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SRC = ROOT / "_diplom_6pz_vps.docx"
DESKTOP_OUT = Path(r"c:\Users\Леонид\Desktop\диплом\6 ПЗ_vps.docx")


def iter_body_blocks(doc: Document):
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield "p", Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield "t", Table(child, doc)


def iter_all_paragraphs(doc: Document):
    for _, block in iter_body_blocks(doc):
        if isinstance(block, Paragraph):
            yield block
        else:
            for row in block.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p


def set_plain_text(paragraph: Paragraph, text: str) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}r"):
            p.remove(child)
    r = paragraph.add_run(text)
    r.bold = False
    r.italic = False


def transform_table_citations(text: str) -> str:
    if not text:
        return text
    # Только абзацы, состоящие из «Продолжение таблицы» без номера (план тестирования)
    if text.strip() == "Продолжение таблицы":
        return "Продолжение таблицы 12"

    pairs: list[tuple[str, str]] = [
        ("Продолжение таблицы 12", "Продолжение таблицы __T17__"),
        ("Продолжение таблицы 9", "Продолжение таблицы 14"),
        ("Продолжение таблицы __T17__", "Продолжение таблицы 17"),
        ("Таблица 15 - Смета затрат на проект", "Таблица 20 – Смета затрат на проект"),
        ("Таблица 15 – Смета затрат на проект", "Таблица 20 – Смета затрат на проект"),
        ("Таблица 14 – Расчет затрат на функционирование", "Таблица 19 – Расчет затрат на функционирование"),
        ("Таблица 13 – Расчет эксплуатационных расходов", "Таблица 18 – Расчет эксплуатационных расходов"),
        (
            "Таблица 12 – Расчёт чистой прибыли организации от разработки программного продукта",
            "Таблица 17 – Расчёт чистой прибыли организации от разработки программного продукта",
        ),
        ("Таблица 11 – Расчет цены реализации программного продукта", "Таблица 16 – Расчет цены реализации программного продукта"),
        ("Таблица 10 – Себестоимость работ по созданию программного продукта", "Таблица 15 – Себестоимость работ по созданию программного продукта"),
        ("Таблица 9 – Расчет расходов на оплату труда с учетом трудозатрат", "Таблица 14 – Расчет расходов на оплату труда с учетом трудозатрат"),
        ("Таблица 4.2 – Сводные результаты тестирования", "Таблица 13 – Сводные результаты тестирования"),
        ("Таблица 4.1 – План тестирования", "Таблица 12 – План тестирования"),
        ("Таблица 3.6 – Календарный план", "Таблица 11 – Календарный план"),
        ("Таблица 3.5 – Сравнение СУБД", "Таблица 10 – Сравнение СУБД"),
        ("Таблица 3.4 – Сравнение ORM и pg", "Таблица 9 – Сравнение ORM и pg"),
        ("Таблица 3.3 – Сравнение инструментов сборки", "Таблица 8 – Сравнение инструментов сборки"),
        ("Таблица 3.2 – Сравнение React, Angular и Vue.js", "Таблица 7 – Сравнение React, Angular и Vue.js"),
        ("Таблица 3.1 – Сравнение JavaScript и TypeScript", "Таблица 6 – Сравнение JavaScript и TypeScript"),
        ("Таблица 2.1 – Сравнение архитектурных паттернов", "Таблица 5 – Сравнение архитектурных паттернов"),
        ("Таблица 1.4 – Основные нефункциональные требования", "Таблица 4 – Основные нефункциональные требования"),
        ("приведена в таблицах 13,14.", "приведена в таблицах 18 и 19."),
        ("систематизированы в таблице 4.2.", "систематизированы в таблице 13."),
        ("представлен в таблице 4.1.", "представлен в таблице 12."),
        ("продолжительность приведены в таблице 3.6.", "продолжительность приведены в таблице 11."),
        ("(таблица 3.5).", "(таблица 10)."),
        ("приведено в таблице 3.4.", "приведено в таблице 9."),
        ("данными таблицы 3.3.", "данными таблицы 8."),
        ("представлено в таблице 3.2.", "представлено в таблице 7."),
        ("приведены в таблице 3.1.", "приведены в таблице 6."),
        ("приведены в таблице 2.1.", "приведены в таблице 5."),
        ("сгруппированы в таблице 1.4.", "сгруппированы в таблице 4."),
        ("перечисленных в таблице 1.3", "перечисленных в таблице 3"),
        ("сведены в таблицу 1.2.", "сведены в таблицу 2."),
        ("Расчёт сметы представлен в таблице 15 .", "Расчёт сметы представлен в таблице 20."),
        ("Расчёт сметы представлен в таблице 15.", "Расчёт сметы представлен в таблице 20."),
        ("приведена в таблицах 13,14", "приведена в таблицах 18 и 19"),
        ("занесём в таблицу 9 .", "занесём в таблицу 14."),
        ("занесём в таблицу 9.", "занесём в таблицу 14."),
        ("сведём в таблицу 10 .", "сведём в таблицу 15."),
        ("сведём в таблицу 10.", "сведём в таблицу 15."),
        ("приведены в таблице 11 .", "приведены в таблице 16."),
        ("приведены в таблице 11.", "приведены в таблице 16."),
        ("сведём в таблицу 12 .", "сведём в таблицу 17."),
        ("сведём в таблицу 12.", "сведём в таблицу 17."),
    ]
    for old, new in pairs:
        if old in text:
            text = text.replace(old, new)
    return text


def apply_paragraph_transforms(doc: Document) -> None:
    for p in iter_all_paragraphs(doc):
        t = p.text
        if not t:
            continue
        new_t = transform_table_citations(t)
        if new_t != t:
            set_plain_text(p, new_t)


def insert_caption_before_table(doc: Document, trigger_substr: str, caption: str) -> bool:
    blocks = list(iter_body_blocks(doc))
    for i, (kind, obj) in enumerate(blocks):
        if kind != "p":
            continue
        if trigger_substr not in obj.text:
            continue
        if i + 1 >= len(blocks) or blocks[i + 1][0] != "t":
            continue
        table = blocks[i + 1][1]
        np = doc.add_paragraph(caption)
        el = np._p
        el.getparent().remove(el)
        table._tbl.addprevious(el)
        return True
    return False


def fill_empty_caption_after_trigger(doc: Document, triggers: tuple[str, ...], caption: str) -> bool:
    """Ищет абзац сразу после абзаца с trigger: если он пустой — задаёт текст подписи."""
    blocks = list(iter_body_blocks(doc))
    for i, (kind, obj) in enumerate(blocks):
        if kind != "p":
            continue
        if not any(tr in obj.text for tr in triggers):
            continue
        if i + 1 >= len(blocks) or blocks[i + 1][0] != "p":
            continue
        nxt = blocks[i + 1][1]
        if nxt.text.strip():
            continue
        set_plain_text(nxt, caption)
        return True
    return False


def main() -> None:
    src = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_SRC
    if not src.is_file():
        raise SystemExit(f"Нет файла: {src}")

    doc = Document(str(src))

    insert_caption_before_table(
        doc,
        "сведены в таблицу 1.2.",
        "Таблица 2 – Пользовательские функции",
    )
    fill_empty_caption_after_trigger(
        doc,
        ("перечисленных в таблице 1.3", "перечисленных в таблице 3"),
        "Таблица 3 – Функционал административной панели",
    )

    apply_paragraph_transforms(doc)

    doc.save(str(src))
    print("OK:", src)

    try:
        DESKTOP_OUT.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy(src, DESKTOP_OUT)
        print("Copied:", DESKTOP_OUT)
    except OSError as e:
        print("Desktop copy skipped:", e)


if __name__ == "__main__":
    main()
