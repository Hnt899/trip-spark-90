# -*- coding: utf-8 -*-
"""
Сквозная нумерация рисунков в ПЗ: Рисунок 1, 2, 3 … вместо Рисунок 2.1, 2.2 … 3.1.
Меняются только подписи «Рисунок …» и ссылки «на рисунке …» / «На рисунке …»,
чтобы не задеть заголовки разделов вида «2.1 Построение…».

Требует: pip install python-docx

Запуск: python scripts/diplom_renumber_figures.py [путь_к.docx]
По умолчанию: _diplom_6pz_vps.docx → перезапись и копия на рабочий стол.
"""
from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SRC = ROOT / "_diplom_6pz_vps.docx"
DESKTOP_OUT = Path(r"c:\Users\Леонид\Desktop\диплом\6 ПЗ_vps.docx")


def iter_body_blocks(doc: Document):
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def iter_all_paragraphs(doc: Document):
    for block in iter_body_blocks(doc):
        if isinstance(block, Paragraph):
            yield block
        else:
            for row in block.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p


def set_plain_text(paragraph: Paragraph, text: str) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}r"):
            p.remove(child)
    paragraph.add_run(text)


def transform_figure_numbering(text: str) -> str:
    # Старый номер → новый (по порядку появления рисунков в документе)
    ordered = [("3.1", "7"), ("2.6", "6"), ("2.5", "5"), ("2.4", "4"), ("2.3", "3"), ("2.2", "2"), ("2.1", "1")]
    for old, new in ordered:
        for dash in ("\u2013", "\u2014", "-", "\u2011"):  # тире в подписях Word
            text = text.replace(f"Рисунок {old} {dash}", f"Рисунок {new} {dash}")
            text = text.replace(f"Рисунок {old}{dash}", f"Рисунок {new} {dash}")
        text = text.replace(f"на рисунке {old}.", f"на рисунке {new}.")
        text = text.replace(f"На рисунке {old}.", f"На рисунке {new}.")
        text = text.replace(f"На рисунке {old} ", f"На рисунке {new} ")
        text = text.replace(f"на рисунке {old} ", f"на рисунке {new} ")
        text = text.replace(f"На рисунке {old}", f"На рисунке {new}")
        text = text.replace(f"на рисунке {old}", f"на рисунке {new}")
    return text


def apply_to_document(doc: Document) -> None:
    for p in iter_all_paragraphs(doc):
        t = p.text
        if not t:
            continue
        new_t = transform_figure_numbering(t)
        if new_t != t:
            set_plain_text(p, new_t)


def verify(doc: Document) -> None:
    bad = []
    for p in iter_all_paragraphs(doc):
        t = p.text
        if re.search(r"Рисунок\s+[23]\.\d", t):
            bad.append(t[:100])
        if re.search(r"[Нн]а\s+рисунке\s+[23]\.\d", t):
            bad.append(t[:100])
    if bad:
        raise SystemExit("Остались старые номера рисунков:\n" + "\n".join(bad))


def main() -> None:
    src = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_SRC
    if not src.is_file():
        raise SystemExit(f"Нет файла: {src}")

    doc = Document(str(src))
    apply_to_document(doc)
    verify(doc)
    doc.save(str(src))
    print("OK:", src)

    try:
        DESKTOP_OUT.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy(src, DESKTOP_OUT)
        print("Copied:", DESKTOP_OUT)
    except OSError as e:
        print("Desktop copy skipped:", e)


if __name__ == "__main__":
    main()
